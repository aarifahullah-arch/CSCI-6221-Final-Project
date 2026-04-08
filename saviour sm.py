# resume_parser.py
# Saviour's Module: Parse different resume formats (PDF, .docx, .doc)

import os
import re
from typing import Dict, Optional

# PDF parsing
try:
    import PyPDF2
except ImportError:
    print("PyPDF2 not installed. Run: pip install PyPDF2")

# DOCX parsing
try:
    import docx
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")

# Legacy .doc parsing (using antiword or textract as fallback)
try:
    import textract
except ImportError:
    print("textract not installed. Run: pip install textract")
    print("Also可能需要安装: sudo apt-get install antiword (Linux) or brew install antiword (Mac)")


class ResumeParser:
    """Parse resumes from different file formats"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']
    
    def parse_resume(self, file_path: str) -> Dict[str, any]:
        """
        Main parsing function that detects file type and extracts text
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Dictionary containing:
                - 'text': extracted text content
                - 'format': file format detected
                - 'success': boolean
                - 'error': error message (if any)
        """
        
        # Check if file exists
        if not os.path.exists(file_path):
            return {
                'text': '',
                'format': None,
                'success': False,
                'error': f"File not found: {file_path}"
            }
        
        # Get file extension
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext not in self.supported_formats:
            return {
                'text': '',
                'format': ext,
                'success': False,
                'error': f"Unsupported format: {ext}. Supported: {self.supported_formats}"
            }
        
        # Parse based on format
        if ext == '.pdf':
            return self._parse_pdf(file_path)
        elif ext == '.docx':
            return self._parse_docx(file_path)
        elif ext == '.doc':
            return self._parse_doc(file_path)
    
    def _parse_pdf(self, file_path: str) -> Dict[str, any]:
        """Parse PDF files using PyPDF2"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from all pages
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
            
            # Clean up the extracted text
            text = self._clean_text(text)
            
            return {
                'text': text,
                'format': 'PDF',
                'success': True,
                'error': None,
                'metadata': {
                    'num_pages': len(pdf_reader.pages)
                }
            }
            
        except Exception as e:
            return {
                'text': '',
                'format': 'PDF',
                'success': False,
                'error': f"PDF parsing error: {str(e)}"
            }
    
    def _parse_docx(self, file_path: str) -> Dict[str, any]:
        """Parse DOCX files using python-docx"""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            
            text = self._clean_text(text)
            
            return {
                'text': text,
                'format': 'DOCX',
                'success': True,
                'error': None,
                'metadata': {
                    'num_paragraphs': len(doc.paragraphs),
                    'num_tables': len(doc.tables)
                }
            }
            
        except Exception as e:
            return {
                'text': '',
                'format': 'DOCX',
                'success': False,
                'error': f"DOCX parsing error: {str(e)}"
            }
    
    def _parse_doc(self, file_path: str) -> Dict[str, any]:
        """Parse legacy .doc files using textract"""
        try:
            # textract works with .doc files
            text = textract.process(file_path).decode('utf-8')
            text = self._clean_text(text)
            
            return {
                'text': text,
                'format': 'DOC',
                'success': True,
                'error': None,
                'metadata': {
                    'method': 'textract'
                }
            }
            
        except Exception as e:
            return {
                'text': '',
                'format': 'DOC',
                'success': False,
                'error': f"DOC parsing error: {str(e)}. Try converting to PDF or DOCX."
            }
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text by removing extra whitespace and normalizing"""
        if not text:
            return ""
        
        # Remove extra newlines and spaces
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Remove special characters that might cause issues
        text = re.sub(r'[^\x00-\x7F]+', ' ', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract email and phone number from resume text (bonus feature)"""
        info = {
            'email': None,
            'phone': None,
            'name': None  # Basic name extraction
        }
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            info['email'] = emails[0]
        
        # Extract phone (US format example)
        phone_pattern = r'\b(?:\+?1[-.]?)?\(?[0-9]{3}\)?[-.]?[0-9]{3}[-.]?[0-9]{4}\b'
        phones = re.findall(phone_pattern, text)
        if phones:
            info['phone'] = phones[0]
        
        # Simple name extraction (first line might contain name)
        lines = text.split('\n')
        for line in lines[:5]:  # Check first 5 lines
            if len(line.split()) <= 4 and len(line) < 50:
                info['name'] = line.strip()
                break
        
        return info


# =====================
# Usage Examples & Test Functions
# =====================

def test_resume_parser():
    """Test function to demonstrate usage"""
    
    parser = ResumeParser()
    
    # Test with different file paths (update these paths)
    test_files = [
        'sample_resume.pdf',
        'sample_resume.docx', 
        'sample_resume.doc'
    ]
    
    for file_path in test_files:
        print(f"\n{'='*50}")
        print(f"Parsing: {file_path}")
        print('='*50)
        
        result = parser.parse_resume(file_path)
        
        if result['success']:
            print(f"✓ Successfully parsed {result['format']} file")
            print(f"  Text length: {len(result['text'])} characters")
            print(f"  Preview: {result['text'][:200]}...")
            
            # Extract contact info
            contact = parser.extract_contact_info(result['text'])
            print(f"  Email: {contact['email']}")
            print(f"  Phone: {contact['phone']}")
            print(f"  Name (guess): {contact['name']}")
            
            if 'metadata' in result:
                print(f"  Metadata: {result['metadata']}")
        else:
            print(f"✗ Failed: {result['error']}")


def parse_resume_for_api(file_path: str) -> dict:
    """
    Simplified function for API integration
    Returns only essential data for backend
    """
    parser = ResumeParser()
    result = parser.parse_resume(file_path)
    
    if result['success']:
        return {
            'status': 'success',
            'resume_text': result['text'],
            'format': result['format'],
            'contact_info': parser.extract_contact_info(result['text'])
        }
    else:
        return {
            'status': 'error',
            'error': result['error']
        }


# =====================
# Quick Setup Instructions
# =====================

